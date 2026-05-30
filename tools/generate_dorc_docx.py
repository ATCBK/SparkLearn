from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Dorc" / "正式交付文档"
PROJECT_START = "2026 年 4 月 9 日"
PROJECT_END = "2026 年 6 月 20 日"


@dataclass
class DocSpec:
    folder: str
    filename: str
    title: str
    audience: str
    standard: str
    purpose: str
    sections: list[tuple[str, list[str]]]


COMMON_TERMS = [
    ("StudyBuddy", "SparkLearn 对本地智能体运行时改造后的自研学习伙伴内核，用于任务执行、工具调用、学习陪伴和多步骤工作流编排。"),
    ("CrewAI", "用于多智能体角色编排、任务分派、协作执行和结果聚合的智能体协作框架。"),
    ("讯飞星辰 Agent 平台", "用于承接平台级智能体能力、资源生成、工具调度和教学任务执行的外部智能体平台。"),
    ("讯飞星火", "平台主要大模型能力来源，承担对话、生成、推理、内容优化等模型调用。"),
    ("RAG", "检索增强生成机制，通过知识库检索、证据片段注入和可信回答控制提升回答准确性。"),
    ("PolarDB PostgreSQL", "生产推荐数据库形态，承接用户画像、资源、对话、记忆、任务和业务事件等结构化数据。"),
    ("JSON/JSONL", "用于画像快照、事件流、Agent 记忆补充层和调试审计日志的轻量数据格式。"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size=10.5, bold=False, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc: Document, text: str = "", style: str | None = None, bold_prefix: bool = False):
    p = doc.add_paragraph(style=style)
    if bold_prefix and "：" in text:
        prefix, rest = text.split("：", 1)
        r1 = p.add_run(prefix + "：")
        set_run_font(r1, bold=True)
        r2 = p.add_run(rest)
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "D9EAF7")
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                for run in p.runs:
                    set_run_font(run)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.3)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for name, size in [("Title", 20), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        styles[name].font.name = "Microsoft YaHei"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True


def add_cover(doc: Document, spec: DocSpec) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SparkLearn 软件工程交付文档")
    set_run_font(r, 18, True, "1F4E79")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(spec.title)
    set_run_font(r, 22, True, "1F4E79")
    doc.add_paragraph()
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["项目名称", "SparkLearn 个性化学习多智能体平台"],
            ["文档版本", "V1.0 正式版"],
            ["项目周期", f"{PROJECT_START} 至 {PROJECT_END}"],
            ["适用对象", spec.audience],
            ["规范依据", spec.standard],
            ["人名处理", "所有示例人员均采用“同学A、教师A、开发者A、管理员A、评审专家A”等占位称呼。"],
        ],
    )
    doc.add_page_break()


def add_revision_and_terms(doc: Document, spec: DocSpec) -> None:
    doc.add_heading("修订记录", level=1)
    add_table(
        doc,
        ["版本", "日期", "修订说明", "修订人"],
        [["V1.0", "2026-05-30", "依据软件工程交付规范形成正式版文档。", "开发者A"]],
    )
    doc.add_heading("文档目的与范围", level=1)
    add_para(doc, f"文档目的：{spec.purpose}", bold_prefix=True)
    add_para(doc, f"适用范围：本文件适用于 SparkLearn 在 {PROJECT_START} 至 {PROJECT_END} 项目周期内形成的设计、开发、测试、部署、演示和交付活动。", bold_prefix=True)
    add_para(doc, f"规范依据：{spec.standard}", bold_prefix=True)
    doc.add_heading("术语与缩略语", level=1)
    add_table(doc, ["术语", "说明"], [[k, v] for k, v in COMMON_TERMS])


def add_standard_sections(doc: Document, spec: DocSpec) -> None:
    for title, paragraphs in spec.sections:
        doc.add_heading(title, level=1)
        for para in paragraphs:
            if para.startswith("- "):
                add_bullet(doc, para[2:])
            elif para.startswith("表："):
                kind = para[2:]
                if kind == "roles":
                    add_table(
                        doc,
                        ["角色", "职责", "典型操作"],
                        [
                            ["同学A", "完成画像采集、学习路径查看、多模态资源学习、练习评测和学习报告查看。", "提交表单、上传资料、发起 AI 对话、生成资源。"],
                            ["教师A", "查看班级学习状态、进行教学干预、分发资源并跟踪学习效果。", "查看大屏、筛选学生、推送资源、查看报告。"],
                            ["管理员A", "维护平台配置、数据库、SDK 凭据、部署环境和运行监控。", "配置 PostgreSQL、维护密钥、处理部署问题。"],
                            ["开发者A", "维护前后端、智能体、数据库、测试和交付文档。", "提交代码、联调 SDK、扩展 Agent 工具。"],
                        ],
                    )
                elif kind == "quality":
                    add_table(
                        doc,
                        ["质量属性", "设计要求", "落地方式"],
                        [
                            ["可靠性", "核心链路应支持异常兜底、错误提示和日志追踪。", "PostgreSQL 持久化、JSONL 事件流、Agent 执行日志、可信回答控制。"],
                            ["可维护性", "模块职责清晰，SDK 接入、业务接口、Agent 编排分层。", "前端页面分层、后端路由分层、StudyBuddy 工具层独立。"],
                            ["可扩展性", "可新增智能体、资源类型、外部工具和数据源。", "CrewAI 编排、StudyBuddy 工具插件、讯飞星辰 Agent 平台能力扩展。"],
                            ["安全性", "密钥不进入文档正文，用户数据按最小必要原则使用。", ".env 管理凭据、数据库权限隔离、接口输入校验。"],
                        ],
                    )
                elif kind == "schedule":
                    add_table(
                        doc,
                        ["阶段", "时间", "主要成果"],
                        [
                            ["启动与需求确认", "2026-04-09 至 2026-04-18", "完成赛题拆解、产品定位、功能树和核心闭环设计。"],
                            ["核心功能开发", "2026-04-19 至 2026-05-12", "完成画像、路径、资源生成、练习、报告、AI 对话等主链路。"],
                            ["智能体与数据深化", "2026-05-13 至 2026-05-30", "完成 CrewAI、StudyBuddy、讯飞星辰 Agent 平台融合，完成 PostgreSQL 迁移和记忆层重构。"],
                            ["测试优化与交付", "2026-05-31 至 2026-06-20", "完成测试验证、部署固化、文档交付、演示材料整理和答辩准备。"],
                        ],
                    )
            else:
                add_para(doc, para)


def save_doc(spec: DocSpec) -> None:
    folder = OUT / spec.folder
    folder.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_document(doc)
    add_cover(doc, spec)
    add_revision_and_terms(doc, spec)
    add_standard_sections(doc, spec)
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("附录：文档关系", level=1)
    add_para(doc, "本文件与 SparkLearn 其他交付文档共同构成软件工程文档集。总体设计类文档说明系统边界和设计原则；技术设计类文档说明前后端与 SDK 接入；智能体与数据架构类文档说明 Agent 协作、数据模型和记忆层；测试部署运维类文档说明质量保障与运行环境；用户使用与协作类文档说明使用流程和团队协作规则。")
    doc.save(folder / spec.filename)


def make_specs() -> list[DocSpec]:
    return [
        DocSpec(
            "01-总体设计类",
            "SparkLearn-UIUX设计说明与系统视觉风格设计.docx",
            "UIUX 设计说明与系统视觉风格设计",
            "评审专家A、产品负责人A、设计师A、前端开发者A",
            "GB/T 8567-2006 计算机软件文档编制规范；ISO/IEC/IEEE 15289 生命周期信息项；用户界面设计一致性原则。",
            "说明 SparkLearn 的界面体系、视觉语言、交互路径、移动端适配规则和系统级视觉风格，确保产品体验与个性化学习场景一致。",
            [
                ("1. 设计目标", [
                    "SparkLearn 的 UIUX 设计围绕“个性化学习闭环”展开，首屏即呈现学习任务、画像状态、路径进度和 AI 辅导入口，避免以营销页替代真实学习工具。",
                    "系统视觉强调清晰、可信、温和和可持续学习，页面以任务推进和数据反馈为核心，不使用过度装饰的卡片堆叠和无意义大面积插画。",
                    "学生端、教师端和管理端采用一致的导航框架、组件语言和反馈机制，确保同学A、教师A、管理员A在跨端使用时形成稳定认知。",
                ]),
                ("2. 信息架构", [
                    "学生端包含首页总览、画像采集、学习路径、资源生成、AI 对话、练习与错题、学习报告、文件与知识库、多模态资源预览等模块。",
                    "教师端包含班级大屏、学生列表、学生详情、教学干预、资料分发、报告查看和课堂反馈等模块。",
                    "移动端适配采用响应式布局，优先保障画像采集、AI 对话、学习路径、资源预览和练习提交等高频任务的可用性。",
                    "表：roles",
                ]),
                ("3. 视觉风格规范", [
                    "主色用于学习进度、按钮状态和关键任务入口；辅助色用于资源类型区分、置信度提示、知识点掌握度和教学干预等级。",
                    "字体层级遵循“页面标题、模块标题、正文说明、辅助标签”四级体系，避免正文过小导致阅读负担。",
                    "组件半径、间距、阴影和边框统一管理。卡片只用于独立信息单元，不把整页区域包装成多层卡片。",
                    "图标优先表达动作含义，例如上传、生成、播放、下载、筛选、搜索、保存、刷新、展开和收起。按钮文字用于明确命令，不用于解释功能原理。",
                ]),
                ("4. 关键流程体验", [
                    "画像采集流程采用问卷、对话和多模态输入结合的方式，同学A可以通过文字、表单和随手拍材料触发个性化逆构。",
                    "学习路径页面以知识阶段、当前节点、推荐资源和掌握度反馈为核心，支持从路径直接跳转到资源生成、练习和 AI 辅导。",
                    "AI 对话支持文本、图片、文件、知识片段和页面上下文，多模态内容在对话区域中直接渲染，减少页面切换。",
                    "视频资源、PPT、文档、思维导图、代码案例和练习题采用统一资源卡片与预览容器，保证资源沉淀和复用。",
                ]),
                ("5. 质量要求", [
                    "表：quality",
                ]),
            ],
        ),
        DocSpec(
            "02-前后端技术设计类",
            "SparkLearn-前端技术设计文档.docx",
            "前端技术设计文档",
            "前端开发者A、测试人员A、评审专家A",
            "IEEE 1016 软件设计说明；GB/T 8567-2006 概要设计说明与详细设计说明。",
            "说明前端架构、页面组织、状态管理、多模态资源技术实现、视频渲染、AI 对话多模态展示和移动端适配方案。",
            [
                ("1. 前端架构概述", [
                    "SparkLearn 前端采用 Next.js App Router 结构组织学生端、教师端、广场端、认证页和多模态资源页面。",
                    "页面层负责业务流程和数据请求，组件层负责复用 UI，API 适配层负责屏蔽真实服务与演示数据之间的差异。",
                    "全局布局包含侧边导航、顶部状态、AI 助手入口和当前用户学习上下文，保证同学A在任意页面都可以进入 AI 辅导。",
                ]),
                ("2. 页面与组件设计", [
                    "核心页面包括首页、画像、路径、资源生成、知识库、练习、报告、视频、Agent 学习伙伴、教师大屏和学生详情。",
                    "组件库包括按钮、输入框、进度条、标签、资源卡片、空状态、错误状态、骨架屏、对话气泡、文件上传器和多模态预览组件。",
                    "所有组件需要具备加载、成功、失败、空数据和权限不足等状态，避免用户在长任务或 SDK 异常时失去反馈。",
                ]),
                ("3. 多模态资源渲染", [
                    "文档资源以 Markdown/富文本预览方式渲染，支持标题层级、代码块、表格、引用和下载入口。",
                    "PPT 资源通过生成结果地址或本地资源记录进行预览，前端提供资源标题、生成状态、页数摘要和打开入口。",
                    "视频资源以任务状态和播放器两层结构呈现，视频生成完成后进入统一资源预览；讲解内容按教学脚本模板输出，避免单一模拟文本。",
                    "图片生成结果以 base64 或资源 URL 方式展示，并在 AI 对话中作为一等消息类型渲染。",
                    "代码案例资源包含讲解、可复制代码、练习任务、扩展要求和评测建议，适配现代编程学习场景。",
                ]),
                ("4. AI 对话前端实现", [
                    "AI 对话通过流式事件更新消息，前端按 text、image、file、source、confidence、done、error 等类型增量渲染。",
                    "对话上下文包含页面来源、用户画像摘要、当前知识点、选中文件和历史消息，用于后端 Agent 编排。",
                    "多模态消息在同一时间轴中呈现，同学A可以上传文件、请求生成图片、查看证据来源并继续追问。",
                ]),
                ("5. 移动端适配", [
                    "移动端采用单列布局、底部优先操作区、可折叠导航和触控友好的按钮尺寸。",
                    "画像采集、路径查看、AI 对话、资源预览和练习提交是移动端首要保障场景。",
                    "表：quality",
                ]),
            ],
        ),
        DocSpec(
            "02-前后端技术设计类",
            "SparkLearn-后端SDK接入与接口设计文档.docx",
            "后端 SDK 接入与接口设计文档",
            "后端开发者A、架构师A、测试人员A、运维人员A",
            "GB/T 8567-2006 概要设计说明、详细设计说明、接口需求说明；ISO/IEC/IEEE 12207 软件生命周期过程。",
            "说明 SparkLearn 后端作为 SDK 接入平台的职责边界、平台能力封装、业务接口组织、异常处理和安全配置。",
            [
                ("1. 后端定位", [
                    "SparkLearn 后端不是单纯的 REST API 服务，而是 SDK 接入平台、智能体编排入口、资源生成调度中心和数据持久化中枢。",
                    "后端统一封装讯飞星火、讯飞 TTS、讯飞智文 PPT、讯飞星辰 Agent 平台、CrewAI 和 StudyBuddy 的接入差异。",
                    "前端只感知业务接口和任务状态，不直接接触 SDK 密钥、外部平台签名、模型调用细节和工具执行过程。",
                ]),
                ("2. SDK 接入层设计", [
                    "讯飞星火 SDK 能力用于智能问答、教学讲解、提示词优化、资源内容生成和个性化逆构。",
                    "讯飞 TTS SDK 能力用于视频讲解、语音合成和多模态资源配音，发音人选择由后端统一校验与映射。",
                    "讯飞智文 PPT 接入用于根据教学大纲、知识点和讲解脚本生成演示材料。",
                    "讯飞星辰 Agent 平台负责外部智能体能力接入、资源任务分发和平台级 Agent 执行。",
                    "CrewAI 与 StudyBuddy 负责本地智能体协作、工具调用、浏览器任务、文件处理和学习伙伴工作流。",
                ]),
                ("3. 业务接口组织", [
                    "画像接口负责问卷、对话、多模态输入和画像逆构结果写入。",
                    "学习路径接口负责根据画像、掌握度和知识图谱关系生成路径节点。",
                    "资源接口负责文档、PPT、视频、图片、代码案例、练习题等多模态资源生成与沉淀。",
                    "对话接口负责 AI Tutor、多模态消息、文件上下文、RAG 证据和防幻觉元数据返回。",
                    "教师端接口负责班级数据、学生画像、干预建议、资料分发和报告聚合。",
                ]),
                ("4. 异常处理与安全", [
                    "所有 SDK 调用必须具备超时、重试、降级和错误日志。外部平台返回失败时，业务接口返回可解释错误，不暴露密钥和内部堆栈。",
                    "密钥通过环境变量管理，部署文档只说明配置项，不记录真实凭据。",
                    "文件上传、资源生成、数据库写入和 Agent 工具调用均需要输入校验，防止非法路径、超大文件和不受控工具执行。",
                    "表：quality",
                ]),
            ],
        ),
        DocSpec(
            "03-智能体与数据架构类",
            "SparkLearn-CrewAI-StudyBuddy-讯飞星辰Agent平台多智能体协作设计文档.docx",
            "CrewAI + StudyBuddy + 讯飞星辰 Agent 平台多智能体协作设计文档",
            "架构师A、算法开发者A、后端开发者A、评审专家A",
            "IEEE 1016 软件设计说明；ISO/IEC/IEEE 15289 生命周期信息项。",
            "说明 SparkLearn 多智能体角色体系、协作机制、任务编排、工具调用、记忆读取和可信输出控制。",
            [
                ("1. 架构目标", [
                    "多智能体架构用于解决个性化学习中的任务复杂度问题，将画像分析、路径规划、资源生成、教学问答、工具调用和质量审核拆分为可观察的角色。",
                    "CrewAI 承担多角色任务编排和协作流程，StudyBuddy 承担本地执行、工具调用和学习伙伴体验，讯飞星辰 Agent 平台承担平台级智能体能力和外部任务执行。",
                    "三者组合形成“本地可控 + 平台增强 + 多角色协作”的混合智能体体系。",
                ]),
                ("2. Agent 角色体系", [
                    "ProfileAgent 负责读取和更新同学A的画像信号，包含目标、水平、薄弱点、偏好、时间约束和当前阶段。",
                    "PlannerAgent 负责将学习目标拆解为路径节点、学习阶段、资源需求和练习安排。",
                    "TutorAgent 负责与同学A进行多模态教学对话，结合页面上下文、RAG 证据和记忆层生成回答。",
                    "ResourceAgent 负责生成文档、PPT、视频、图片、练习题、代码案例和拓展阅读等资源。",
                    "ToolAgent 负责浏览器工具、文件工具、外部信息源、资料分发和平台工具调用。",
                    "ReviewerAgent 负责防幻觉检查、证据一致性判断、输出格式审核和风险提示。",
                ]),
                ("3. 协作流程", [
                    "请求进入系统后先由 IntentAgent 判断任务类型，再由编排层决定是否调用画像、RAG、工具、资源生成和审核链路。",
                    "对话型任务采用“意图识别 -> 画像读取 -> RAG 检索 -> Tutor 生成 -> Reviewer 审核 -> 记忆更新”的流程。",
                    "资源型任务采用“目标解析 -> 模板选择 -> 内容生成 -> 多模态渲染 -> 资源沉淀 -> 学习路径挂接”的流程。",
                    "教学干预任务采用“教师A选择对象 -> 班级数据聚合 -> 薄弱点识别 -> 资源匹配 -> 分发记录 -> 效果追踪”的流程。",
                ]),
                ("4. 可观测性", [
                    "每次 Agent 执行应记录任务 ID、角色名称、输入摘要、工具调用、状态变化、耗时、输出摘要和错误信息。",
                    "执行日志进入 PostgreSQL 和 JSONL 事件流，支持调试、回放、答辩展示和质量追踪。",
                    "表：quality",
                ]),
            ],
        ),
        DocSpec(
            "03-智能体与数据架构类",
            "SparkLearn-CrewAI-StudyBuddy-星辰Agent混合架构与PostgreSQL数据库设计文档.docx",
            "CrewAI + StudyBuddy + 星辰 Agent 混合架构与 PostgreSQL 数据库设计文档",
            "架构师A、数据库开发者A、后端开发者A、运维人员A",
            "GB/T 8567-2006 数据要求说明、概要设计说明；ISO/IEC/IEEE 12207。",
            "说明 SparkLearn 的混合智能体架构、PostgreSQL 主库设计、JSON/JSONL 辅助数据层、用户画像表单和资源沉淀模型。",
            [
                ("1. 总体架构", [
                    "系统采用前端应用层、后端 SDK 接入层、智能体编排层、数据持久化层和外部平台接入层五层结构。",
                    "PostgreSQL 是核心业务主库，推荐生产环境使用阿里云 PolarDB PostgreSQL，以满足国产云环境、弹性扩展和关系型查询需求。",
                    "JSON/JSONL 用于画像快照、事件审计、Agent 中间结果、离线调试和快速回放，不替代 PostgreSQL 的核心持久化地位。",
                ]),
                ("2. 数据库主题域", [
                    "用户与身份域：记录用户基础信息、角色、权限、登录状态和班级归属。",
                    "画像域：记录问卷表单、对话采集、多模态逆构结果、画像版本、画像快照和字段来源。",
                    "学习路径域：记录阶段、知识点、路径节点、节点状态、推荐资源和掌握度关联。",
                    "资源域：记录文档、PPT、视频、图片、代码案例、练习题、外部资源和资源模板。",
                    "对话与记忆域：记录会话、消息、多模态附件、RAG 证据、Agent 记忆、执行日志和更新事件。",
                    "教师端域：记录班级、学生概况、教学干预、资料分发、报告聚合和大屏指标。",
                ]),
                ("3. 用户画像与表单模型", [
                    "画像表单包含学习目标、当前水平、薄弱知识点、学习偏好、可用时间、实践能力、认知风格和当前阶段。",
                    "多模态逆构能力支持同学A上传随手拍、资料片段或自由描述，由模型反推学习需求、项目目标、阶段规划和资源建议。",
                    "画像采用版本化管理，每次问卷提交、对话沉淀、练习结果和手动编辑都会形成可追踪更新。",
                ]),
                ("4. JSON/JSONL 辅助层", [
                    "profile_snapshot.json 保存当前画像快照，便于快速注入上下文和离线演示。",
                    "learning_events.jsonl 保存学习事件流，包括画像更新、资源生成、练习提交、报告生成和 Agent 调用。",
                    "memory_store.json 保存 Agent 记忆补充层，辅助长期偏好、约束、事实、弱点和阶段信息沉淀。",
                    "表：quality",
                ]),
            ],
        ),
        DocSpec(
            "03-智能体与数据架构类",
            "SparkLearn-数据流转与Agent记忆层设计文档.docx",
            "数据流转与 Agent 记忆层设计文档",
            "架构师A、后端开发者A、算法开发者A、测试人员A",
            "GB/T 8567-2006 详细设计说明；ISO/IEC/IEEE 15289。",
            "说明用户数据、学习行为、多模态资源、RAG 证据、Agent 记忆和防幻觉结果在系统中的流转方式。",
            [
                ("1. 数据流转总览", [
                    "SparkLearn 数据流从同学A的输入开始，经过画像采集、路径规划、资源生成、对话辅导、练习评测、报告反馈和画像更新形成闭环。",
                    "每个关键节点都会产生结构化业务数据和事件日志，前者进入 PostgreSQL，后者进入 JSONL 审计流。",
                    "Agent 记忆层读取画像、历史行为、对话摘要、掌握度和文件上下文，为后续任务提供个性化依据。",
                ]),
                ("2. 画像采集流", [
                    "问卷流：同学A提交目标、水平、薄弱点、偏好和时间信息，系统生成初始画像。",
                    "对话流：同学A通过自然语言描述学习需求，系统抽取目标、约束、偏好和阶段。",
                    "多模态逆构流：同学A上传图片、资料或项目想法，系统反推学习规划、技能缺口和资源建议。",
                    "手动编辑流：同学A或教师A对画像字段进行确认与修正，系统保留版本变化。",
                ]),
                ("3. RAG 与防幻觉流", [
                    "系统先根据问题、页面上下文、用户画像和选中文件构造检索查询，再从知识库、上传文件和资源库中召回证据片段。",
                    "ReviewerAgent 根据证据数量、相关性、来源可信度和回答风险生成置信度元数据。",
                    "最终回答必须带有证据边界和可解释来源，避免无依据扩展，保证教学场景下的可信输出。",
                ]),
                ("4. 记忆层设计", [
                    "Working Memory 保存当前会话上下文和临时任务状态。",
                    "Episodic Memory 保存学习事件、对话片段、资源使用和任务执行历史。",
                    "Semantic Memory 保存长期稳定画像，如目标、偏好、弱点、约束、技能和学习阶段。",
                    "Perceptual Memory 保存文件、图片、随手拍和多模态输入解析出的可复用信息。",
                    "表：quality",
                ]),
            ],
        ),
        DocSpec(
            "03-智能体与数据架构类",
            "SparkLearn-个性化画像采集文档.docx",
            "个性化画像采集文档",
            "产品负责人A、后端开发者A、算法开发者A、评审专家A",
            "GB/T 8567-2006 数据要求说明；ISO/IEC/IEEE 15289。",
            "说明 SparkLearn 如何通过表单、对话、多模态输入和学习行为持续采集、构建和更新个性化画像。",
            [
                ("1. 画像目标", [
                    "画像系统的目标是让平台理解同学A的学习目标、基础水平、薄弱点、学习方式、时间约束和当前阶段。",
                    "画像不是一次性档案，而是随学习行为持续更新的动态模型。",
                    "画像直接驱动路径规划、资源生成、AI 辅导、练习推荐、报告总结和教师端干预。",
                ]),
                ("2. 采集方式", [
                    "表单采集：通过多步问卷采集目标、水平、薄弱点、偏好和时间。",
                    "对话采集：通过 AI 引导式问答补充学习背景、项目目标和约束条件。",
                    "多模态采集：通过随手拍、资料上传、代码片段、错题图片等输入反推学习需求。",
                    "行为采集：通过资源使用、练习结果、错题记录、对话主题和路径推进记录更新画像。",
                ]),
                ("3. 字段规范", [
                    "核心字段包括 user_id、goal、knowledge_level、weak_points、learning_preference、daily_time、cognitive_style、practical_ability、current_stage、version 和 updated_at。",
                    "字段来源需要记录 source_type，例如 onboarding、dialog、file_upload、practice、teacher_edit、agent_inference。",
                    "画像更新应记录变更前后值、触发事件、置信度和更新时间，便于审计和回滚。",
                ]),
                ("4. 应用场景", [
                    "路径规划根据画像生成阶段路径和知识点顺序。",
                    "资源生成根据画像选择模板、难度、讲解方式和练习类型。",
                    "AI 对话根据画像调整解释粒度、示例风格和学习建议。",
                    "教师端根据画像聚合班级薄弱点并生成教学干预建议。",
                ]),
                ("5. 质量控制", [
                    "表：quality",
                ]),
            ],
        ),
        DocSpec(
            "04-测试部署运维类",
            "SparkLearn-测试文档.docx",
            "测试文档",
            "测试人员A、开发者A、评审专家A、运维人员A",
            "ISO/IEC/IEEE 29119 软件测试文档；GB/T 8567-2006 测试计划、测试分析报告。",
            "说明 SparkLearn 的测试范围、测试策略、测试环境、测试用例设计、质量标准和验收结论。",
            [
                ("1. 测试范围", [
                    "测试范围覆盖学生端、教师端、后端 SDK 接入、PostgreSQL 数据库、多智能体编排、StudyBuddy 工具调用、讯飞星辰 Agent 平台接入、多模态资源生成和部署运行。",
                    "核心业务链路包括画像采集、路径规划、资源生成、AI 对话、文件上传、RAG 检索、防幻觉、练习评测、学习报告和教师干预。",
                    "表：schedule",
                ]),
                ("2. 测试策略", [
                    "单元测试验证后端路由、数据模型、工具函数和核心业务逻辑。",
                    "集成测试验证前后端联调、SDK 封装、PostgreSQL 读写、Agent 编排和文件处理。",
                    "端到端测试验证同学A从画像建立到资源生成、练习评测和报告查看的完整闭环。",
                    "兼容性测试验证桌面端、移动端、主流浏览器和不同屏幕尺寸下的显示与交互。",
                ]),
                ("3. 重点测试用例", [
                    "画像采集用例：提交表单、对话补充、多模态逆构、画像版本更新和路径刷新。",
                    "多模态资源用例：生成文档、PPT、视频、图片、代码案例和练习题，并验证资源沉淀。",
                    "AI 对话用例：上传文件后提问，验证 RAG 证据、图片生成、置信度元数据和记忆更新。",
                    "教师端用例：教师A查看班级数据、选择同学A、分发资源并查看学习报告。",
                ]),
                ("4. 质量结论", [
                    "系统按照软件工程测试分层完成核心链路验证，满足演示、答辩和项目交付要求。",
                    "表：quality",
                ]),
            ],
        ),
        DocSpec(
            "04-测试部署运维类",
            "SparkLearn-部署文档.docx",
            "部署文档",
            "运维人员A、后端开发者A、前端开发者A、管理员A",
            "GB/T 8567-2006 安装计划、操作手册；ISO/IEC/IEEE 12207 运维过程。",
            "说明 SparkLearn 的环境要求、数据库配置、SDK 凭据配置、前后端启动、生产部署和故障排查。",
            [
                ("1. 环境要求", [
                    "操作系统建议使用 Windows 10/11 或云服务器 Linux 环境，开发演示环境可使用本地 Windows。",
                    "后端运行环境为 Python 3.12 及以上，前端运行环境为 Node.js 20 LTS 及以上。",
                    "数据库使用 PostgreSQL，生产环境推荐 PolarDB PostgreSQL。",
                    "外部能力包括讯飞星火、讯飞 TTS、讯飞智文 PPT、讯飞星辰 Agent 平台、CrewAI 和 StudyBuddy。",
                ]),
                ("2. 配置项", [
                    "数据库配置包含 PostgreSQL host、port、database、user、password、sslmode 和连接池参数。",
                    "SDK 配置包含讯飞星火、TTS、智文 PPT、讯飞星辰 Agent 平台等凭据，所有密钥写入环境变量，不写入代码和文档正文。",
                    "文件存储配置包含本地上传目录、资源生成目录、静态资源访问路径和可选 OSS 存储参数。",
                ]),
                ("3. 启动流程", [
                    "后端启动前需要安装 Python 依赖、初始化数据库表、检查环境变量和 SDK 连通性。",
                    "前端启动前需要安装 npm 依赖、配置后端接口地址并执行构建检查。",
                    "StudyBuddy 启动前需要确认工作区、工具权限、浏览器自动化依赖和任务执行目录。",
                    "表：schedule",
                ]),
                ("4. 故障排查", [
                    "数据库连接失败时检查网络、安全组、账号权限、SSL 设置和迁移脚本执行状态。",
                    "SDK 调用失败时检查密钥、额度、模型名称、接口地址、签名参数和超时设置。",
                    "资源生成失败时检查任务日志、Agent 执行步骤、文件写入权限和模板渲染结果。",
                    "表：quality",
                ]),
            ],
        ),
        DocSpec(
            "05-用户使用与协作类",
            "SparkLearn-项目使用手册.docx",
            "项目使用手册",
            "同学A、教师A、管理员A、评审专家A",
            "GB/T 8567-2006 用户手册、操作手册。",
            "面向普通使用者、教师、管理员和评审专家说明 SparkLearn 的主要功能、操作路径和演示流程。",
            [
                ("1. 产品概述", [
                    "SparkLearn 是面向个性化学习的多智能体平台，提供画像采集、学习路径、资源生成、AI 辅导、练习评测、学习报告和教师干预等功能。",
                    "项目周期为 2026 年 4 月 9 日至 2026 年 6 月 20 日，交付目标是形成可演示、可测试、可部署和可扩展的软件系统。",
                    "表：roles",
                ]),
                ("2. 学生端使用流程", [
                    "同学A进入平台后先完成画像采集，可通过问卷、对话或多模态输入建立学习档案。",
                    "同学A进入学习路径页面查看当前阶段、知识点节点、推荐资源和后续学习建议。",
                    "同学A可以生成文档、PPT、视频、图片、代码案例和练习题等多模态资源，并在资源库中沉淀复用。",
                    "同学A可以在任意页面打开 AI 对话，上传文件、提问、生成图片、查看证据来源并继续追问。",
                    "同学A完成练习后查看判题结果、错题记录、掌握度变化和学习报告。",
                ]),
                ("3. 教师端使用流程", [
                    "教师A进入教师端后查看班级总览、学生分布、薄弱知识点、学习活跃度和报告摘要。",
                    "教师A可以打开学生详情，查看同学A的画像、路径、练习结果和干预建议。",
                    "教师A可以选择资料分发对象，向指定学生或班级推送学习资源。",
                ]),
                ("4. 管理与演示流程", [
                    "管理员A负责配置数据库、SDK 凭据、资源目录和运行环境。",
                    "评审专家A可按“画像 -> 路径 -> 资源 -> 对话 -> 练习 -> 报告 -> 教师端”的顺序体验完整闭环。",
                    "表：quality",
                ]),
            ],
        ),
        DocSpec(
            "05-用户使用与协作类",
            "SparkLearn-开发协作指导文档.docx",
            "开发协作指导文档",
            "开发者A、测试人员A、产品负责人A、管理员A",
            "GB/T 8567-2006 软件开发计划、维护手册；ISO/IEC/IEEE 12207 软件生命周期过程。",
            "说明 SparkLearn 项目开发协作、任务分工、分支管理、联调流程、文档维护和质量检查规则。",
            [
                ("1. 协作目标", [
                    "开发协作以“功能可演示、代码可维护、数据可追踪、文档可交付”为目标。",
                    "项目周期为 2026 年 4 月 9 日至 2026 年 6 月 20 日，开发者A、测试人员A、产品负责人A和管理员A按照阶段目标推进。",
                    "表：schedule",
                ]),
                ("2. 分工规则", [
                    "前端开发者A负责页面、组件、多模态渲染、移动端适配和用户交互。",
                    "后端开发者A负责 SDK 接入、业务接口、数据库、文件服务、RAG、防幻觉和资源任务。",
                    "算法开发者A负责 CrewAI、StudyBuddy、讯飞星辰 Agent 平台协作、记忆层和工具扩展。",
                    "测试人员A负责测试计划、测试用例、缺陷记录、回归验证和验收报告。",
                ]),
                ("3. 联调规范", [
                    "接口联调前需要确认请求字段、响应字段、错误码、流式事件类型和数据落库结果。",
                    "多模态资源联调需要同时验证生成任务、前端渲染、资源沉淀和再次打开效果。",
                    "Agent 联调需要保留任务 ID、执行步骤、工具调用和最终输出，便于定位问题。",
                ]),
                ("4. 文档维护", [
                    "设计变更必须同步更新对应交付文档，避免代码能力和材料描述不一致。",
                    "所有文档中的示例人员统一使用同学A、教师A、管理员A、开发者A和评审专家A。",
                    "表：quality",
                ]),
            ],
        ),
        DocSpec(
            "00-文档编制说明",
            "SparkLearn-软件工程文档编制说明.docx",
            "软件工程文档编制说明",
            "项目负责人A、开发者A、评审专家A",
            "GB/T 8567-2006；ISO/IEC/IEEE 12207；ISO/IEC/IEEE 15289；IEEE 1016；ISO/IEC/IEEE 29119。",
            "说明本次 SparkLearn 交付文档的目录、命名、格式、状态口径、人名替换规则和软件工程规范依据。",
            [
                ("1. 文档集范围", [
                    "本次文档集覆盖总体设计、前端技术、后端 SDK 接入、智能体协作、数据库设计、数据流转、画像采集、测试、部署、使用手册和开发协作。",
                    "文档统一存放在 Dorc/正式交付文档 目录下，并按照文档类型划分子目录。",
                    "项目周期统一写作 2026 年 4 月 9 日至 2026 年 6 月 20 日。",
                ]),
                ("2. 编写规范", [
                    "每份文档包含封面信息、修订记录、文档目的、适用范围、术语、主体章节、质量要求和附录。",
                    "文档内容采用已实现技术方案口径描述 CrewAI、StudyBuddy、讯飞星辰 Agent 平台、PostgreSQL、多模态资源、RAG、记忆层和防幻觉系统。",
                    "文档不出现真实人名，所有角色示例使用同学A、教师A、开发者A、管理员A、评审专家A等占位称呼。",
                ]),
                ("3. 目录规范", [
                    "00-文档编制说明：保存文档编写规范与文档关系。",
                    "01-总体设计类：保存系统总体、UIUX、视觉风格等文档。",
                    "02-前后端技术设计类：保存前端、后端 SDK 接入与接口设计文档。",
                    "03-智能体与数据架构类：保存多智能体、数据库、数据流转、画像采集等文档。",
                    "04-测试部署运维类：保存测试文档和部署文档。",
                    "05-用户使用与协作类：保存项目使用手册和开发协作指导文档。",
                ]),
                ("4. 质量要求", [
                    "表：quality",
                ]),
            ],
        ),
    ]


def main() -> None:
    for spec in make_specs():
        save_doc(spec)
    print(f"Generated {len(make_specs())} docx files under {OUT}")


if __name__ == "__main__":
    main()
